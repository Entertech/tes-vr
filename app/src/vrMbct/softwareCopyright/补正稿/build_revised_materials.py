from __future__ import annotations

import copy
import re
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[5]
MATERIAL_DIR = ROOT / "app/src/vrMbct/softwareCopyright"
OUTPUT_DIR = MATERIAL_DIR / "补正稿"
SCREENSHOT_DIR = MATERIAL_DIR / "screenshots"

FULL_NAME = "VR-MBCT正念认知神经调控训练APP"
SHORT_NAME = "VR-MBCT"
VERSION = "V1.0"
HEADER_TEXT = f"{FULL_NAME} {VERSION}"
EAST_BODY = "Noto Sans CJK SC"
EAST_HEAD = "Noto Sans CJK SC"


def set_run_font(run, east_asia: str, latin: str, size: float, bold: bool = False):
    run.font.name = east_asia
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), east_asia)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), east_asia)
    run._element.rPr.rFonts.set(qn("w:cs"), east_asia)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, value, end):
        run._r.append(node)
    set_run_font(run, EAST_BODY, "Times New Roman", 8)


def add_bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)


def configure_page(section):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    section.header_distance = Mm(7)
    section.footer_distance = Mm(7)


def configure_header(section):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.clear()
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(8.7), WD_TAB_ALIGNMENT.CENTER)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(17.4), WD_TAB_ALIGNMENT.RIGHT)
    paragraph.add_run("\t")
    title_run = paragraph.add_run(HEADER_TEXT)
    set_run_font(title_run, EAST_BODY, "Times New Roman", 8)
    paragraph.add_run("\t")
    add_page_field(paragraph)
    add_bottom_border(paragraph)


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = EAST_BODY
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_BODY)
    normal._element.rPr.rFonts.set(qn("w:ascii"), EAST_BODY)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), EAST_BODY)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.45
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.first_line_indent = Cm(0.74)

    for name, size in (("Heading 1", 15), ("Heading 2", 12.5)):
        style = styles[name]
        style.font.name = EAST_HEAD
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_HEAD)
        style._element.rPr.rFonts.set(qn("w:ascii"), EAST_HEAD)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), EAST_HEAD)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(7)
        style.paragraph_format.keep_with_next = True

    if "Figure Caption" not in styles:
        caption = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Figure Caption"]
    caption.font.name = EAST_BODY
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_BODY)
    caption._element.rPr.rFonts.set(qn("w:ascii"), EAST_BODY)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), EAST_BODY)
    caption.font.size = Pt(9)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(0)
    caption.paragraph_format.first_line_indent = Cm(0)


def add_body(document, text: str, *, indent=True, bold_lead: str | None = None):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0.74) if indent else Cm(0)
    paragraph.paragraph_format.line_spacing = 1.45
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, EAST_HEAD, "Arial", 10.5, True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest, EAST_BODY, "Times New Roman", 10.5)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, EAST_BODY, "Times New Roman", 10.5)
    return paragraph


def add_step(document, number: int, text: str):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.75)
    paragraph.paragraph_format.first_line_indent = Cm(-0.6)
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.paragraph_format.space_after = Pt(3)
    marker = paragraph.add_run(f"{number}. ")
    set_run_font(marker, EAST_BODY, "Times New Roman", 10, True)
    run = paragraph.add_run(text)
    set_run_font(run, EAST_BODY, "Times New Roman", 10)
    return paragraph


def add_bullet(document, text: str):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.8)
    paragraph.paragraph_format.first_line_indent = Cm(-0.55)
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.paragraph_format.space_after = Pt(3)
    marker = paragraph.add_run("• ")
    set_run_font(marker, EAST_BODY, "Times New Roman", 10.5)
    run = paragraph.add_run(text)
    set_run_font(run, EAST_BODY, "Times New Roman", 10.5)
    return paragraph


def add_figure(document, filename: str, caption: str, width_cm: float = 7.2):
    image_path = SCREENSHOT_DIR / filename
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = document.add_paragraph(caption, style="Figure Caption")
    return cap


def add_page_break(document):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def build_manual():
    document = Document()
    section = document.sections[0]
    configure_page(section)
    configure_header(section)
    configure_styles(document)

    # Page 1
    document.add_heading("1. 系统介绍", level=1)
    add_body(document, f"{FULL_NAME}（简称：{SHORT_NAME}）用于在 Android 终端上配合 TES 设备执行正念认知训练流程。系统将账号管理、设备连接、训练前引导、课程选择、刺激过程控制和会话记录组织为连续操作链路。")
    add_body(document, "软件采用 Kotlin 与 XML 开发，通过 AndroidX 组件组织页面，通过蓝牙低功耗通信能力连接设备。训练过程中的会话编号、阶段状态、时间戳、课程信息和设备消息以结构化记录保存在应用私有目录中。")
    document.add_heading("1.1 运行环境", level=2)
    add_bullet(document, "运行平台：Android 7.0 及以上版本。")
    add_bullet(document, "硬件条件：支持 BLE 的 Android 终端及配套 TES 设备。")
    add_bullet(document, "通信条件：启用蓝牙，并授予系统要求的蓝牙连接、扫描及通知权限。")
    document.add_heading("1.2 标准操作流程", level=2)
    for idx, text in enumerate((
        "启动软件并进入账号入口。",
        "登录已有训练账号，或登记新的本地训练账号。",
        "连接 TES 设备并核对设备状态。",
        "进入 VR-MBCT 模式，完成训练前引导。",
        "选择冥想课程并启动训练。",
        "完成训练后保存本次会话记录。",
    ), 1):
        add_step(document, idx, text)
    add_page_break(document)

    # Page 2
    document.add_heading("2. 软件启动", level=1)
    add_body(document, "在 Android 终端点击软件图标后，系统加载应用资源并进入 VR-MBCT 账号入口。首次运行时，Android 系统可能显示蓝牙、附近设备或通知权限申请，应根据设备连接需要完成授权。")
    add_step(document, 1, "点击软件图标，等待启动画面结束。")
    add_step(document, 2, "确认蓝牙已打开，并按系统提示授予所需权限。")
    add_step(document, 3, "进入账号入口后继续登录或注册。")
    add_figure(document, "02_login.png", "图 1 软件启动界面")
    add_page_break(document)

    # Page 3
    document.add_heading("3. 账号登录与注册", level=1)
    add_body(document, "已有账号的使用者输入账号和密码后进入主界面。首次使用时，可进入注册页面登记姓名、账号、手机号、机构或科室及密码。账号信息保存在应用本地，用于识别当前训练使用者。")
    add_step(document, 1, "依次填写姓名、账号、手机号和机构信息。")
    add_step(document, 2, "设置密码并再次输入相同密码。")
    add_step(document, 3, "点击【完成注册】，系统保存账号并进入主界面。")
    add_figure(document, "03_register.png", "图 2 训练账号注册界面")
    add_page_break(document)

    # Page 4
    document.add_heading("4. 主界面", level=1)
    add_body(document, "主界面显示当前账号、所属机构、课程库数量和本地记录数量，并集中提供设备连接、模式入口、训练数据和账号管理入口。开始训练时，建议首先进入设备连接页面。")
    add_step(document, 1, "核对欢迎区域中的姓名、账号和机构信息。")
    add_step(document, 2, "查看课程库数量及本地记录数量。")
    add_step(document, 3, "点击【连接设备并开始训练】进入设备连接流程。")
    add_figure(document, "04_home.png", "图 3 系统主界面")
    add_page_break(document)

    # Page 5
    document.add_heading("5. 连接训练设备", level=1)
    add_body(document, "设备连接页面支持通过 MAC 地址或设备名称发起连接。已知固定设备地址时宜使用 MAC 地址连接；仅掌握设备名称时可使用名称连接。连接过程中应保持设备开机并处于可连接状态。")
    add_step(document, 1, "输入目标设备的 MAC 地址或设备名称。")
    add_step(document, 2, "点击对应的连接按钮，等待连接结果。")
    add_step(document, 3, "状态显示为已连接后进入模式页面。")
    add_figure(document, "05_connect.png", "图 4 设备连接界面")
    add_page_break(document)

    # Page 6
    document.add_heading("6. 查看设备信息", level=1)
    add_body(document, "设备信息页面用于核对当前连接状态、设备名称、MAC 地址、设备运行信息和最近设备消息。连接状态变化后，可点击刷新按钮重新读取当前信息。")
    add_step(document, 1, "进入设备信息页面。")
    add_step(document, 2, "点击【刷新设备信息】。")
    add_step(document, 3, "确认连接状态、名称与地址和实际设备一致。")
    add_figure(document, "06_device_info.png", "图 5 设备信息界面")
    add_page_break(document)

    # Page 7
    document.add_heading("7. 进入 VR-MBCT 模式", level=1)
    add_body(document, "设备连接完成后进入模式入口页面。选择 VR-MBCT 训练入口，系统创建训练流程并转入训练前引导页面。需要再次核对设备时，可先打开设备信息页面。")
    add_step(document, 1, "确认设备已连接。")
    add_step(document, 2, "点击 VR-MBCT 训练入口。")
    add_step(document, 3, "等待系统进入训练前引导页面。")
    add_figure(document, "07_mode_entry.png", "图 6 VR-MBCT 模式入口")
    add_page_break(document)

    # Page 8
    document.add_heading("8. 训练前引导", level=1)
    add_body(document, "训练前引导阶段用于建立会话、显示倒计时并记录设备状态。页面同步显示训练阶段、剩余时间、设备运行信息以及脑波曲线区域。倒计时结束后，系统进入课程选择页面。")
    add_step(document, 1, "保持设备佩戴与连接状态稳定。")
    add_step(document, 2, "按照页面提示完成三分钟训练前引导。")
    add_step(document, 3, "倒计时结束后等待课程列表自动打开。")
    add_figure(document, "08_prepare.png", "图 7 训练前引导界面")
    add_page_break(document)

    # Page 9
    document.add_heading("9. 选择训练课程", level=1)
    add_body(document, "课程页面以卡片形式显示可选的 MBCT 冥想课程，每张卡片包含课程名称、训练时长和简要说明。选择课程后，系统记录课程编号、课程名称和时长，并进入疗程执行页面。")
    add_step(document, 1, "浏览课程名称与训练说明。")
    add_step(document, 2, "点击目标课程卡片中的【选择该课程】。")
    add_step(document, 3, "在疗程执行页面核对课程名称和初始时长。")
    add_figure(document, "09_course_list.png", "图 8 MBCT 课程选择界面")
    add_page_break(document)

    # Page 10
    document.add_heading("10. 开始疗程", level=1)
    add_body(document, "疗程执行页面显示已选课程、刺激状态、引导状态、倒计时、会话状态、设备信息和本地数据文件位置。开始前应核对课程名称、设备状态和倒计时是否正确。")
    add_step(document, 1, "确认已选课程和显示时长。")
    add_step(document, 2, "确认设备保持连接且状态正常。")
    add_step(document, 3, "点击【开始刺激】进入训练过程。")
    add_figure(document, "10_session_initial.png", "图 9 疗程待开始界面")
    add_page_break(document)

    # Page 11
    document.add_heading("11. 训练过程监控", level=1)
    add_body(document, "训练启动后，页面持续更新刺激阶段、剩余时间、会话状态和曲线区域。训练状态按准备、启动、执行、结束及末次引导顺序变化。除需要提前终止外，不应在训练过程中关闭页面或断开蓝牙。")
    add_step(document, 1, "观察刺激状态和剩余时间。")
    add_step(document, 2, "保持终端与设备在有效连接范围内。")
    add_step(document, 3, "如需提前结束，点击【停止刺激】并等待停止结果。")
    add_figure(document, "11_session_running.png", "图 10 疗程运行界面")
    add_page_break(document)

    # Page 12
    document.add_heading("12. 末次引导与会话完成", level=1)
    add_body(document, "课程阶段结束后，系统进入末次引导阶段，并继续维护当前会话。末次引导结束时，系统写入会话完成节点并保留本次数据文件路径。")
    document.add_heading("12.1 完成条件", level=2)
    add_bullet(document, "课程刺激状态已结束。")
    add_bullet(document, "末次引导倒计时已完成。")
    add_bullet(document, "页面提示本次会话完成，数据文件已保存。")
    document.add_heading("12.2 提前停止", level=2)
    add_body(document, "使用者点击停止按钮后，软件向设备发送停止请求，并记录停止操作及设备返回结果。停止完成前应保持蓝牙连接，不应立即退出软件。")
    document.add_heading("12.3 返回主界面", level=2)
    add_body(document, "会话完成后可关闭当前页面并返回主界面。再次训练时，系统会创建新的会话编号和独立记录文件，不覆盖之前的训练记录。")
    add_page_break(document)

    # Page 13
    document.add_heading("13. 账号中心", level=1)
    add_body(document, "账号中心显示当前使用者的姓名、账号、手机号、机构和登录状态。使用者可退出当前账号、重新登记账号或返回主界面。退出登录不会删除已经保存的训练记录。")
    add_step(document, 1, "核对当前账号与机构信息。")
    add_step(document, 2, "需要更换使用者时点击【退出登录】。")
    add_step(document, 3, "需要更新本地账号时点击【重新注册账号】。")
    add_figure(document, "15_account.png", "图 11 账号中心界面")
    add_page_break(document)

    # Page 14
    document.add_heading("14. 训练记录保存", level=1)
    add_body(document, "每次训练使用独立会话编号。过程记录保存到 Android 应用外部私有目录的 vr_mbct 文件夹，文件名采用 mbct_session_<会话编号>.jsonl 格式。每行保存一个阶段事件，便于按发生顺序读取。")
    document.add_heading("14.1 主要记录字段", level=2)
    for text in (
        "sessionId：标识一次独立训练会话。",
        "stage：标识当前事件所属的训练阶段。",
        "timestamp：记录事件发生时间。",
        "courseId、courseTitle、duration：记录已选课程及训练时长。",
        "deviceState、deviceMessage：记录关键节点的设备状态与消息。",
    ):
        add_bullet(document, text)
    document.add_heading("14.2 主要阶段节点", level=2)
    add_body(document, "训练记录覆盖训练前引导开始与完成、课程选择、刺激请求、刺激开始、刺激完成、末次引导开始与完成、手动停止、设备断连和会话完成等节点。")
    document.add_heading("14.3 数据保护", level=2)
    add_body(document, "账号和训练记录保存在应用本地。使用者应妥善管理终端访问权限，不应通过非授权方式复制包含个人信息的记录文件。")
    add_page_break(document)

    # Page 15
    document.add_heading("15. 常见操作与异常处理", level=1)
    document.add_heading("15.1 无法连接设备", level=2)
    add_body(document, "确认设备已开机、蓝牙已开启，并检查附近设备及蓝牙权限。通过 MAC 地址连接时，应核对地址格式；通过名称连接时，应核对设备名称并缩短终端与设备的距离。")
    document.add_heading("15.2 训练过程中连接中断", level=2)
    add_body(document, "保持当前页面，检查设备供电和蓝牙状态。软件会记录设备断连节点。恢复连接前不应继续发起新的刺激流程。")
    document.add_heading("15.3 无法开始训练", level=2)
    add_body(document, "依次确认账号已登录、设备已连接、训练前引导已经完成且课程已经选择。若页面状态未更新，可返回上一页重新进入，但不要同时启动多个训练页面。")
    document.add_heading("15.4 账号信息错误", level=2)
    add_body(document, "进入账号中心退出当前账号，再通过注册入口重新登记正确的姓名、账号、手机号和机构信息。")
    document.add_heading("15.5 本地记录核对", level=2)
    add_body(document, "会话完成后检查页面显示的数据文件路径。不同会话使用不同文件名；终端清理应用数据或卸载软件前，应按管理要求备份需要保留的训练记录。")

    out = OUTPUT_DIR / f"{FULL_NAME}_{VERSION}_操作说明书_补正稿.docx"
    document.save(out)
    return out


SOURCE_FILES = [
    "app/src/vrMbct/java/com/entertech/tes/vr/mode/mbct/MbctModels.kt",
    "app/src/vrMbct/java/com/entertech/tes/vr/mode/mbct/MbctRecordStore.kt",
    "app/src/vrMbct/java/com/entertech/tes/vr/mode/mbct/MbctUserStore.kt",
    "app/src/vrMbct/java/com/entertech/tes/vr/mode/mbct/MbctBrainwaveSupport.kt",
    "app/src/vrMbct/java/com/entertech/tes/vr/mode/mbct/MbctBrainwaveChartView.kt",
    "app/src/vrMbct/java/com/entertech/tes/vr/mode/mbct/MbctPrepareViewModel.kt",
    "app/src/vrMbct/java/com/entertech/tes/vr/mode/mbct/MbctPrepareActivity.kt",
    "app/src/vrMbct/java/com/entertech/tes/vr/mode/mbct/MbctCourseAdapter.kt",
    "app/src/vrMbct/java/com/entertech/tes/vr/mode/mbct/MbctCourseListActivity.kt",
    "app/src/vrMbct/java/com/entertech/tes/vr/mode/mbct/MbctSessionViewModel.kt",
    "app/src/vrMbct/java/com/entertech/tes/vr/mode/mbct/MbctSessionActivity.kt",
    "app/src/main/java/com/entertech/tes/vr/connect/ConnectDeviceActivity.kt",
    "app/src/vrMbct/java/com/entertech/tes/vr/mode/mbct/MbctDeviceInfoViewModel.kt",
    "app/src/vrMbct/java/com/entertech/tes/vr/mode/mbct/MbctDeviceInfoActivity.kt",
    "app/src/vrMbct/java/com/entertech/tes/vr/mode/mbct/MbctLoginActivity.kt",
    "app/src/vrMbct/java/com/entertech/tes/vr/mode/mbct/MbctRegisterActivity.kt",
    "app/src/vrMbct/java/com/entertech/tes/vr/mode/mbct/MbctAccountActivity.kt",
    "app/src/vrMbct/AndroidManifest.xml",
    "app/src/vrMbct/res/layout/activity_mbct_account.xml",
    "app/src/vrMbct/res/layout/activity_mbct_course_list.xml",
    "app/src/vrMbct/res/layout/activity_mbct_device_info.xml",
    "app/src/vrMbct/res/layout/activity_mbct_login.xml",
    "app/src/vrMbct/res/layout/activity_mbct_prepare.xml",
    "app/src/vrMbct/res/layout/activity_mbct_register.xml",
    "app/src/vrMbct/res/layout/activity_mbct_session.xml",
    "app/src/vrMbct/res/layout/choose_mode_activity.xml",
    "app/src/vrMbct/res/layout/item_mbct_course.xml",
    "app/src/vrMbct/res/values/colors.xml",
    "app/src/vrMbct/res/values/strings.xml",
]


def source_display_lines(width=116):
    displayed = []
    for rel in SOURCE_FILES:
        path = ROOT / rel
        if not path.exists():
            raise FileNotFoundError(path)
        raw_lines = path.read_text(encoding="utf-8").expandtabs(4).splitlines()
        for line in raw_lines:
            if len(line) <= width:
                displayed.append(line)
                continue
            leading = re.match(r"\s*", line).group(0)
            wrapped = textwrap.wrap(
                line,
                width=width,
                subsequent_indent=leading + "    ",
                replace_whitespace=False,
                drop_whitespace=False,
                break_long_words=False,
                break_on_hyphens=False,
            )
            displayed.extend(wrapped or [line])
    if len(displayed) < 3000:
        raise ValueError(f"代表性源码显示行不足 3000 行：{len(displayed)}")
    return displayed[:3000]


def build_source():
    lines = source_display_lines()
    document = Document()
    section = document.sections[0]
    configure_page(section)
    section.top_margin = Mm(16)
    section.bottom_margin = Mm(13)
    configure_header(section)

    normal = document.styles["Normal"]
    normal.font.name = EAST_BODY
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_BODY)
    normal._element.rPr.rFonts.set(qn("w:ascii"), EAST_BODY)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), EAST_BODY)
    normal.font.size = Pt(7)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(10.2)

    for page in range(60):
        page_lines = lines[page * 50:(page + 1) * 50]
        if page_lines and page_lines[-1] == "":
            for index in range(len(page_lines) - 2, -1, -1):
                if page_lines[index] != "":
                    page_lines[index], page_lines[-1] = page_lines[-1], page_lines[index]
                    break
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(10.2)
        run = paragraph.add_run("\n".join(page_lines))
        set_run_font(run, EAST_BODY, "Courier New", 7)
        if page != 59:
            paragraph.add_run().add_break(WD_BREAK.PAGE)

    out = OUTPUT_DIR / f"{FULL_NAME}_{VERSION}_源程序鉴别材料_补正稿.docx"
    document.save(out)
    return out


def replace_cell_text(cell, old: str | None, new: str):
    full = "\n".join(p.text for p in cell.paragraphs)
    if old is not None and old not in full:
        raise ValueError(f"未在单元格中找到待替换内容：{old}")
    result = full.replace(old, new) if old is not None else new
    first = cell.paragraphs[0]
    if first.runs:
        first.runs[0].text = result
        for run in first.runs[1:]:
            run.text = ""
    else:
        first.add_run(result)
    for paragraph in cell.paragraphs[1:]:
        for run in paragraph.runs:
            run.text = ""


def build_info_form():
    src = MATERIAL_DIR / "VR_MBCT_有材料信息采集表-2026版.docx"
    document = Document(src)
    table = document.tables[0]
    replace_cell_text(table.rows[2].cells[2], "VR-MBCT正念认知神经调控训练软件", FULL_NAME)
    replace_cell_text(table.rows[4].cells[2], "例如：V1.0", VERSION)
    replace_cell_text(table.rows[14].cells[2], "约1598行", "约4307行")
    feature_cell = table.rows[18].cells[2]
    replace_cell_text(feature_cell, "VR-MBCT正念认知神经调控训练软件", FULL_NAME)
    for table_item in document.tables:
        for row in table_item.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if re.search(r"[\u3400-\u9fff]", run.text):
                            fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
                            for slot in ("eastAsia", "ascii", "hAnsi", "cs"):
                                fonts.set(qn(f"w:{slot}"), EAST_BODY)
                            run.font.name = EAST_BODY
    # 原模板在“主要功能”首个运行上设置了字符边框，会随自动换行形成密集方框。
    for paragraph in feature_cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)
            r_pr = run._r.get_or_add_rPr()
            border = r_pr.find(qn("w:bdr"))
            if border is not None:
                r_pr.remove(border)
    # 通知书对应单一申请人；保留两行填写位置，去除造成纯空白第三页的冗余行。
    for row_index in range(28, 23, -1):
        row = table.rows[row_index]
        table._tbl.remove(row._tr)
    out = OUTPUT_DIR / "VR_MBCT_材料信息采集表-2026版_补正稿.docx"
    document.save(out)
    return out


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    outputs = [build_manual(), build_source(), build_info_form()]
    for path in outputs:
        print(path)


if __name__ == "__main__":
    main()
